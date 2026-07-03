from fpdf import FPDF
from docx import Document

# ---------- 1) A small multi-page PDF -------------------------------------
pdf = FPDF()
pdf.set_auto_page_break(auto=True, margin=15)
pdf.add_page()
pdf.set_font("Helvetica", "B", 16)
pdf.cell(text="Orbital AI -- RAG System Architecture (v2.1)")
pdf.ln(10)
pdf.set_font("Helvetica", "", 11)

# NOTE: Helvetica is a built-in PDF font that only supports latin-1.
# For unicode (em-dash, CJK, emoji) you must register a TTF (DejaVu etc.).
# Lesson: real-world PDF loaders need to handle this; PyPDFLoader does.
pages = [
    "Section 1 -- Overview\n"
    "Orbital AI's retrieval-augmented generation (RAG) system, codenamed "
    "'Prometheus', is the internal knowledge layer that powers our support "
    "chatbot, the engineering wiki assistant, and the leadership briefing "
    "generator. It is built on a microservices architecture deployed across "
    "two regions (Bengaluru and Frankfurt) for redundancy.\n\n"
    "The system is built on Python 3.11, LangChain 0.3, PostgreSQL 16 with "
    "the pgvector extension, and Google Gemini 2.0 Flash as the primary LLM.",
    "Section 2 -- Indexing Pipeline\n"
    "Documents enter the system through three ingestion paths:\n"
    "  1) Scheduled crawls of internal Confluence spaces (every 6 hours).\n"
    "  2) Webhook-triggered sync of the GitHub monorepo (on push to main).\n"
    "  3) Manual upload portal used by the operations team.\n\n"
    "All documents are normalised to Markdown, chunked with a recursive "
    "character splitter (chunk_size=800, overlap=150), embedded with the "
    "BGE-small-en-v1.5 model, and stored in pgvector with HNSW indexing.",
    "Section 3 -- Retrieval\n"
    "At query time the system performs:\n"
    "  * Query rewriting (HyDE + step-back)\n"
    "  * Hybrid retrieval (BM25 + dense, fused with reciprocal rank)\n"
    "  * Cross-encoder re-ranking (BGE-reranker-base)\n"
    "  * Contextual compression to fit the prompt window\n"
    "Average end-to-end latency: 480 ms p50, 1.2 s p95.",
    "Section 4 -- Safety & Evaluation\n"
    "Every response is evaluated by the offline RAGAS pipeline on a "
    "weekly-rotating golden set of 250 questions. Faithfulness, answer "
    "relevancy, and context precision/recall are tracked in Grafana. "
    "PII redaction happens at ingestion and again at generation.",
]
for content in pages:
    pdf.multi_cell(0, 6, text=content)
    pdf.add_page()

# remove the trailing blank page that add_page at the end of the loop adds
pdf.output("/workspace/rag_course/02_loaders/data/sample.pdf")
print("wrote sample.pdf (4 pages)")

# ---------- 2) A small DOCX -----------------------------------------------
doc = Document()
doc.add_heading("Orbital AI — Engineering Onboarding Guide", level=1)

doc.add_heading("Welcome", level=2)
doc.add_paragraph(
    "Welcome to Orbital AI! This guide covers the first two weeks of your "
    "onboarding. Your buddy will pair with you for the first three days."
)

doc.add_heading("Week 1 — Tooling", level=2)
doc.add_paragraph("Set up the following tools on day 1:", style="List Bullet")
for item in [
    "Git + GitHub CLI (access via SSO)",
    "uv for Python dependency management",
    "Docker Desktop with the Orbital AI image registry",
    "Linear for issue tracking",
    "Slack — channels: #eng, #eng-incidents, #random",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Week 2 — Codebase Tour", level=2)
doc.add_paragraph(
    "Your manager will assign a tour guide. The recommended reading order "
    "for new engineers is: services/core, services/rag, services/api, "
    "infra/terraform. Each module has a README at the root."
)

doc.add_heading("Contact", level=2)
doc.add_paragraph("Reach out to people-success@orbitalai.example for help.")

doc.save("/workspace/rag_course/02_loaders/data/sample.docx")
print("wrote sample.docx")

# ---------- 3) A CSV -------------------------------------------------------
csv_text = """name,role,team,years_at_company
Priya Raman,CEO,Leadership,4
Marcus Holloway,CTO,Engineering,4
Aisha Patel,Head of Research,Research,3
Tomohiro Sato,VP Engineering,Engineering,2
Sofia Lindgren,Senior SRE,Platform,3
Diego Alvarez,RAG Engineer,AI,1
Hannah O'Connor,Security Lead,Security,2
Kenji Tanaka,Product Manager,Product,3
"""
with open("/workspace/rag_course/02_loaders/data/employees.csv", "w") as f:
    f.write(csv_text)
print("wrote employees.csv")

# ---------- 4) A Markdown file ---------------------------------------------
md_text = """# Orbital AI — ML Model Card

## Model: `star-track-reranker-v3`

### Intended use
Cross-encoder re-ranker used inside the Prometheus RAG pipeline to
re-order the top-50 retrieved candidates down to the top-5 that
actually go into the LLM prompt.

### Architecture
`BAAI/bge-reranker-base` fine-tuned on 38,000 internal (query, doc,
relevance) triples mined from production logs.

### Training data
- 38k triples from the Bengaluru RAG query logs (Mar–Dec 2024)
- 12k triples from the Frankfurt logs (Oct–Dec 2024)
- 4k synthetic adversarial triples from the safety team

### Metrics
| Metric             | Value |
|--------------------|-------|
| NDCG@10            | 0.83  |
| MRR                | 0.79  |
| Latency p50        | 38 ms |
| Latency p95        | 91 ms |

### Known limitations
- Degrades on queries longer than 64 tokens (truncation).
- Not yet evaluated on multilingual queries (FR, JA, HI).
- The re-ranker can over-fit to recent query distributions;
  we re-train on a rolling 90-day window.
"""
with open("/workspace/rag_course/02_loaders/data/model_card.md", "w") as f:
    f.write(md_text)
print("wrote model_card.md")

print("\nDone. You can now run: python loaders_demo.py")
